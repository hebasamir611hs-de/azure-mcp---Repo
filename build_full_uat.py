# -*- coding: utf-8 -*-
"""
Build the FULL Camel Race Sprint-1 UAT document from all_cases.json by cloning
the approved template (قالب_حالات_الاختبار_UAT.docx).

Layout: section "حالات الاختبار" -> for each backlog item (PBI), a heading with
the PBI title, then its UAT case tables (numbered UAT-001..N across the whole
document). PBIs with no UAT cases get a heading + an Arabic "not applicable" note
so every backlog item is accounted for (traceability).
"""
import copy
import json
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# reuse helpers from the sample builder
from build_uat_from_template import (
    W, _set_run_text, set_tc_text, set_paragraph_text, para_text,
)

CONFIG = {
    "template":  "/Users/asmaa/CI-CD-AI-AU/قالب_حالات_الاختبار_UAT.docx",
    "cases":     "/Users/asmaa/CI-CD-AI-AU/azure-mcp---Repo/outputs/UAT/_cases/all_cases.json",
    "manifest":  "/Users/asmaa/CI-CD-AI-AU/azure-mcp---Repo/outputs/UAT/_src/_manifest.tsv",
    "output":    "/Users/asmaa/CI-CD-AI-AU/azure-mcp---Repo/outputs/UAT/"
                 "UAT_CamelRace_Sprint1.docx",
    "project_name": "سباقات الهجن",
    "phase":        "تطوير تطبيق الهاتف المحمول – Sprint 1",
    "client":       "اللجنة المنظمة لسباقات الهجن",
    "client_result_label": "نتيجة اختبار العميل",
    "author":       "iHorizons – فريق ضمان الجودة",
    "doc_date":     "2026/06/21",
}


def build_case_table(template_tbl, uat_no, case, client_result_label):
    """Clone the per-case table and fill it for one UAT case (dict form)."""
    tbl = copy.deepcopy(template_tbl)
    trs = tbl.findall(W('w:tr'))
    title_tr, precond_tr, header_tr = trs[0], trs[1], trs[2]
    step_template = trs[3]

    set_tc_text(title_tr.find(W('w:tc')), f"UAT-{uat_no:03d} - {case['title']}")
    set_tc_text(precond_tr.find(W('w:tc')),
                f"المتطلبات الأساسية: {case['preconditions']}")
    if client_result_label:
        set_tc_text(header_tr.findall(W('w:tc'))[4], client_result_label)

    for tr in trs[3:]:
        tbl.remove(tr)
    for i, step in enumerate(case['steps'], start=1):
        tr = copy.deepcopy(step_template)
        cells = tr.findall(W('w:tc'))
        set_tc_text(cells[0], str(i))
        set_tc_text(cells[1], step['action'])
        set_tc_text(cells[2], step['expected'])
        set_tc_text(cells[3], "")
        set_tc_text(cells[4], "")
        tbl.append(tr)
    return tbl


def load_titles(path):
    titles = {}
    with open(path, encoding="utf-8") as f:
        next(f)  # header
        for line in f:
            parts = line.rstrip("\n").split("\t")
            if len(parts) >= 4:
                titles[int(parts[0])] = parts[3]
    return titles


def main():
    cases_path = sys.argv[1] if len(sys.argv) > 1 else CONFIG["cases"]
    output_path = sys.argv[2] if len(sys.argv) > 2 else CONFIG["output"]

    doc = Document(CONFIG["template"])
    data = json.load(open(cases_path, encoding="utf-8"))
    titles = load_titles(CONFIG["manifest"])

    # header / intro replacements (NOT the group-name placeholder; we remove it)
    for p in doc.paragraphs:
        full = para_text(p)
        new = full
        for k, v in {
            "[اسم المشروع]": CONFIG["project_name"],
            "[المرحلة]": CONFIG["phase"],
            "وزارة الأوقاف والشؤون الإسلامية": CONFIG["client"],
            "[اسم المشروع / المرحلة]": f"{CONFIG['project_name']} – {CONFIG['phase']}",
        }.items():
            if k in new:
                new = new.replace(k, v)
        if new != full:
            set_paragraph_text(p, new)

    # version-control table row 1
    vt = doc.tables[1]
    vt.rows[1].cells[0].text = "1.0"
    vt.rows[1].cells[1].text = CONFIG["doc_date"]
    vt.rows[1].cells[2].text = CONFIG["author"]
    vt.rows[1].cells[3].text = "إنشاء الملف"

    # locate template elements
    group_label_p = None
    for p in doc.paragraphs:
        if "[اسم مجموعة حالات الاختبار]" in para_text(p):
            group_label_p = p._p
            break
    normal_p = None
    for p in doc.paragraphs:
        if "[أضف هنا أي وصف" in para_text(p):
            normal_p = p._p
            break

    note_tbl  = doc.tables[2]._tbl   # editor note
    example_a = doc.tables[3]._tbl
    example_b = doc.tables[4]._tbl
    case_tmpl = doc.tables[5]._tbl

    heading_tmpl = copy.deepcopy(group_label_p)   # PBI group heading style
    note_tmpl    = copy.deepcopy(normal_p)        # normal paragraph style

    uat_no = 0
    pbis_with_cases = 0
    for entry in data:
        pbi = entry["pbi"]
        title = titles.get(pbi, "")
        cases = entry.get("cases", [])

        # PBI heading
        h = copy.deepcopy(heading_tmpl)
        set_paragraph_text(h, f"{title}  —  (PBI {pbi})")
        note_tbl.addprevious(h)

        if not cases:
            n = copy.deepcopy(note_tmpl)
            reason = entry.get("note", "").strip() or "لا تنطبق حالات اختبار قبول على هذا البند."
            set_paragraph_text(n, f"لا تنطبق حالات اختبار قبول المستخدم — {reason}")
            note_tbl.addprevious(n)
            continue

        pbis_with_cases += 1
        for case in cases:
            uat_no += 1
            tbl = build_case_table(case_tmpl, uat_no, case, CONFIG["client_result_label"])
            note_tbl.addprevious(tbl)
            note_tbl.addprevious(OxmlElement('w:p'))  # spacer

    # remove template scaffolding
    for el in (group_label_p, note_tbl, example_a, example_b, case_tmpl):
        el.getparent().remove(el)

    doc.save(output_path)
    print("Saved:", output_path)
    print(f"PBIs: {len(data)} | with cases: {pbis_with_cases} | total UAT cases: {uat_no}")


if __name__ == "__main__":
    main()

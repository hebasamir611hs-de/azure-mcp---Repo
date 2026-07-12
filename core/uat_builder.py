"""
core/uat_builder.py — Skill 11b: UAT Document Generation from a Sprint.

Builds a client-ready User Acceptance Testing (UAT) .docx straight from every
backlog item (PBI) in a sprint. Each PBI becomes a UAT section: business
requirement + acceptance criteria + an execution table the testers fill in
(Actual Result / Status / Comments).

This is the mechanical counterpart to the `drafter` subagent / `build-uat-doc`
skill: those build a curated doc from *signed-off, UAT-tagged test cases*; this
tool produces a first-pass template directly from raw sprint requirements, with
no test-case derivation. Use it to bootstrap a UAT pack for a whole sprint.

All functions are plain — no @mcp.tool() decorators. Registration happens in server.py.
"""

import html
import os
import re
from datetime import datetime

from azure.devops.v7_1.work_item_tracking.models import Wiql
from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor

from core.test_planner import _parse_iteration_path
from core.utils import (
    apply_rtl,
    get_azure_client,
    handle_error,
    is_arabic,
    sanitize_wiql_string,
    set_cell_shading,
)


# ─────────────────────────────────────────────────────────────────────────────
# STATIC LABELS (bilingual)
# ─────────────────────────────────────────────────────────────────────────────

_LABELS = {
    "en": {
        "doc_title": "User Acceptance Testing (UAT) Document",
        "doc_control": "Document Control",
        "field": "Field",
        "value": "Value",
        "project": "Project",
        "sprint": "Sprint",
        "generated_on": "Generated On",
        "total_items": "Total Backlog Items",
        "prepared_by": "Prepared By",
        "prepared_by_value": "QA Team (auto-generated)",
        "intro_heading": "1. Introduction",
        "intro_body": (
            "This document defines the User Acceptance Testing (UAT) scenarios for the "
            "backlog items delivered in this sprint. For each item, the business requirement "
            "and acceptance criteria are listed, followed by a verification table for the "
            "business stakeholders to execute and record results."
        ),
        "scope_heading": "2. Scope",
        "scope_body": "The following backlog items are in scope for this UAT cycle:",
        "scope_id": "PBI ID",
        "scope_title": "Title",
        "scope_state": "State",
        "scenarios_heading": "3. UAT Scenarios by Backlog Item",
        "requirement": "Business Requirement",
        "acceptance_criteria": "Acceptance Criteria",
        "no_ac": "No acceptance criteria defined — to be confirmed with the Business Analyst.",
        "no_desc": "No description provided.",
        "col_num": "#",
        "col_scenario": "Scenario",
        "col_steps": "Steps to Verify",
        "col_expected": "Expected Result",
        "col_actual": "Actual Result",
        "col_status": "Status (Pass/Fail)",
        "col_comments": "Comments",
        "scenario_label": "Verify acceptance criterion",
        "step_placeholder": "Execute the flow described by the criterion and observe the outcome.",
        "signoff_heading": "4. UAT Sign-off",
        "signoff_body": (
            "By signing below, the stakeholders confirm that the scenarios above have been "
            "executed and the results are accepted."
        ),
        "role": "Role",
        "name": "Name",
        "signature": "Signature",
        "date": "Date",
        "roles": ["UAT Lead", "Business Owner", "QA Manager", "Project Manager"],
    },
    "ar": {
        "doc_title": "وثيقة اختبار قبول المستخدم (UAT)",
        "doc_control": "ضبط الوثيقة",
        "field": "الحقل",
        "value": "القيمة",
        "project": "المشروع",
        "sprint": "السبرنت",
        "generated_on": "تاريخ الإنشاء",
        "total_items": "إجمالي عناصر العمل",
        "prepared_by": "إعداد",
        "prepared_by_value": "فريق ضمان الجودة (إنشاء تلقائي)",
        "intro_heading": "١. مقدمة",
        "intro_body": (
            "تحدد هذه الوثيقة سيناريوهات اختبار قبول المستخدم (UAT) لعناصر العمل المُسلَّمة في "
            "هذا السبرنت. لكل عنصر، يُدرَج المتطلب والمعايير المقبولة، يليهما جدول تحقق يقوم "
            "أصحاب المصلحة بتنفيذه وتسجيل النتائج فيه."
        ),
        "scope_heading": "٢. النطاق",
        "scope_body": "عناصر العمل التالية مشمولة في دورة اختبار القبول هذه:",
        "scope_id": "رقم العنصر",
        "scope_title": "العنوان",
        "scope_state": "الحالة",
        "scenarios_heading": "٣. سيناريوهات الاختبار حسب عنصر العمل",
        "requirement": "متطلب العمل",
        "acceptance_criteria": "معايير القبول",
        "no_ac": "لا توجد معايير قبول محددة — يجب التأكيد مع محلل الأعمال.",
        "no_desc": "لا يوجد وصف.",
        "col_num": "م",
        "col_scenario": "السيناريو",
        "col_steps": "خطوات التحقق",
        "col_expected": "النتيجة المتوقعة",
        "col_actual": "النتيجة الفعلية",
        "col_status": "الحالة (نجاح/فشل)",
        "col_comments": "ملاحظات",
        "scenario_label": "التحقق من معيار القبول",
        "step_placeholder": "نفّذ المسار الموصوف في المعيار ولاحظ النتيجة.",
        "signoff_heading": "٤. اعتماد اختبار القبول",
        "signoff_body": (
            "بالتوقيع أدناه، يؤكد أصحاب المصلحة أن السيناريوهات أعلاه قد نُفِّذت وأن النتائج "
            "مقبولة."
        ),
        "role": "الدور",
        "name": "الاسم",
        "signature": "التوقيع",
        "date": "التاريخ",
        "roles": ["قائد اختبار القبول", "مالك العمل", "مدير ضمان الجودة", "مدير المشروع"],
    },
}

_HEADER_FILL = "1F4E79"   # dark blue header row
_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)


# ─────────────────────────────────────────────────────────────────────────────
# HTML → TEXT HELPERS (Azure stores Description / AC as HTML)
# ─────────────────────────────────────────────────────────────────────────────

def _html_to_lines(raw: str) -> list:
    """Converts an HTML fragment to a list of plain-text lines, preserving list/paragraph breaks."""
    if not raw:
        return []
    text = re.sub(r'(?i)</(li|p|div|h[1-6]|tr)>', '\n', raw)
    text = re.sub(r'(?i)<br\s*/?>', '\n', text)
    text = re.sub(r'(?i)<li[^>]*>', '• ', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = html.unescape(text)
    lines = [re.sub(r'\s+', ' ', ln).strip(' •\t') for ln in text.split('\n')]
    return [ln for ln in lines if ln]


def _html_to_text(raw: str) -> str:
    """Flattens an HTML fragment to a single whitespace-normalized string."""
    lines = _html_to_lines(raw)
    return " ".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX BUILDING HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _set_table_rtl(table) -> None:
    """Marks a table as right-to-left (visual column order) for Arabic content."""
    tblPr = table._tbl.tblPr
    tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))


def _style_header_row(table, rtl: bool) -> None:
    """Applies the dark-blue fill + white bold text to a table's first row."""
    for cell in table.rows[0].cells:
        set_cell_shading(cell, _HEADER_FILL)
        for para in cell.paragraphs:
            if rtl:
                apply_rtl(para)
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = _HEADER_TEXT


def _add_kv_table(doc, rows: list, labels: dict, rtl: bool) -> None:
    """Adds a two-column Field/Value table from a list of (key, value) tuples."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT if rtl else WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = labels["field"]
    hdr[1].text = labels["value"]
    _style_header_row(table, rtl)
    for key, val in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(val)
    if rtl:
        _set_table_rtl(table)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    apply_rtl(para)


def _add_para(doc, text: str, rtl: bool, bold: bool = False, size: int = None):
    """Adds a paragraph with optional RTL, bold, and font size."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if rtl:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        apply_rtl(para)
    return para


def _add_heading(doc, text: str, level: int, rtl: bool):
    """Adds a heading and applies RTL alignment when needed."""
    heading = doc.add_heading(text, level=level)
    if rtl:
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        apply_rtl(heading)
    return heading


def _build_pbi_section(doc, labels: dict, pbi: dict, doc_rtl: bool) -> int:
    """
    Renders one PBI's UAT section: requirement, acceptance criteria, execution table.
    Returns the number of scenario rows generated for this PBI.
    """
    # Per-PBI language can differ from the document default
    rtl = is_arabic(pbi["title"]) or is_arabic(pbi["ac_text"]) or doc_rtl
    lbl = _LABELS["ar"] if rtl else labels

    _add_heading(doc, f"PBI {pbi['id']} — {pbi['title']}", level=2, rtl=rtl)

    # Business requirement
    _add_para(doc, lbl["requirement"], rtl, bold=True)
    desc = _html_to_text(pbi["description"]) or lbl["no_desc"]
    _add_para(doc, desc, rtl)

    # Acceptance criteria (bulleted)
    _add_para(doc, lbl["acceptance_criteria"], rtl, bold=True)
    ac_lines = _html_to_lines(pbi["ac_raw"])
    if ac_lines:
        for line in ac_lines:
            bullet = doc.add_paragraph(line, style="List Bullet")
            if rtl:
                bullet.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                apply_rtl(bullet)
    else:
        _add_para(doc, lbl["no_ac"], rtl)

    # UAT execution table
    headers = [
        lbl["col_num"], lbl["col_scenario"], lbl["col_steps"],
        lbl["col_expected"], lbl["col_actual"], lbl["col_status"], lbl["col_comments"],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT if rtl else WD_TABLE_ALIGNMENT.LEFT
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    _style_header_row(table, rtl)

    scenario_rows = ac_lines if ac_lines else [lbl["no_ac"]]
    for idx, criterion in enumerate(scenario_rows, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = f"{lbl['scenario_label']} #{idx}"
        cells[2].text = lbl["step_placeholder"]
        cells[3].text = criterion           # the criterion IS the concrete expected outcome
        cells[4].text = ""                  # Actual — tester fills
        cells[5].text = ""                  # Status — tester fills
        cells[6].text = ""                  # Comments — tester fills

    if rtl:
        _set_table_rtl(table)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    apply_rtl(para)

    return len(scenario_rows)


# ─────────────────────────────────────────────────────────────────────────────
# SKILL 11b: BUILD UAT DOCUMENT FROM A SPRINT
# ─────────────────────────────────────────────────────────────────────────────

def create_uat_document_from_sprint(
    sprint_input: str,
    output_path: str = "",
    language: str = "auto",
) -> dict:
    """
    SKILL 11b: Generate a UAT Document for an Entire Sprint

    Reads every backlog item (PBI) in the given sprint and builds a single,
    client-ready User Acceptance Testing .docx. Each PBI becomes a UAT section
    containing its business requirement, acceptance criteria, and an execution
    table (Scenario / Steps / Expected / Actual / Status / Comments) for the
    business stakeholders to fill in.

    This is a *template generator from raw requirements* — it does NOT derive or
    require injected test cases. For a curated doc built from signed-off,
    UAT-tagged test cases, use the build-uat-doc skill / drafter subagent instead.

    Args:
        sprint_input (str): Iteration path (e.g. 'Woqod\\MCP-test-3') OR the full
                            sprint URL copied from the Azure DevOps browser.
        output_path (str): Optional .docx destination path. When omitted, the file
                           is written to '<repo>/outputs/UAT/UAT_<sprint>_<timestamp>.docx'.
        language (str): 'auto' (default — detect from PBI content), 'en', or 'ar'.
                        Controls the document's static section labels; per-PBI
                        content keeps its own detected direction regardless.

    Returns:
        {status, file_path, sprint, language, total_pbis, total_scenarios, pbis}
    """
    try:
        iteration_path = _parse_iteration_path(sprint_input)
        project = os.getenv("AZURE_PROJECT")
        client = get_azure_client()

        safe_iter = sanitize_wiql_string(iteration_path)
        safe_proj = sanitize_wiql_string(project)

        query = f"""
            SELECT [System.Id], [System.Title], [System.State],
                   [Microsoft.VSTS.Common.AcceptanceCriteria], [System.Description]
            FROM WorkItems
            WHERE [System.WorkItemType] = 'Product Backlog Item'
              AND [System.IterationPath] = '{safe_iter}'
              AND [System.TeamProject] = '{safe_proj}'
              AND [System.State] <> 'Removed'
            ORDER BY [System.Id] ASC
        """
        result = client.query_by_wiql(Wiql(query=query))

        if not result.work_items:
            return {
                "status": "no_pbis",
                "sprint": iteration_path,
                "total_pbis": 0,
                "message": f"No backlog items found in sprint '{iteration_path}'.",
            }

        ids = [wi.id for wi in result.work_items]
        work_items = client.get_work_items(ids=ids, expand="All")

        pbis = []
        for item in work_items:
            if item is None:
                continue
            f = item.fields
            ac_raw = f.get("Microsoft.VSTS.Common.AcceptanceCriteria", "") or ""
            pbis.append({
                "id": item.id,
                "title": f.get("System.Title", f"PBI {item.id}"),
                "state": f.get("System.State", ""),
                "description": f.get("System.Description", "") or "",
                "ac_raw": ac_raw,
                "ac_text": _html_to_text(ac_raw),
            })

        # Resolve document-level language
        if language not in ("en", "ar"):
            corpus = " ".join(p["title"] + " " + p["ac_text"] for p in pbis)
            language = "ar" if is_arabic(corpus) else "en"
        labels = _LABELS[language]
        doc_rtl = language == "ar"

        # ── Build the document ──────────────────────────────────────────────
        doc = DocxDocument()
        sprint_short = iteration_path.split("\\")[-1]
        generated_on = datetime.now().strftime("%Y-%m-%d %H:%M")

        title = doc.add_heading(labels["doc_title"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = _add_para(doc, f"{project}  —  {sprint_short}", doc_rtl, bold=True, size=14)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _add_heading(doc, labels["doc_control"], level=1, rtl=doc_rtl)
        _add_kv_table(doc, [
            (labels["project"], project),
            (labels["sprint"], iteration_path),
            (labels["generated_on"], generated_on),
            (labels["total_items"], len(pbis)),
            (labels["prepared_by"], labels["prepared_by_value"]),
        ], labels, doc_rtl)

        _add_heading(doc, labels["intro_heading"], level=1, rtl=doc_rtl)
        _add_para(doc, labels["intro_body"], doc_rtl)

        # Scope table
        _add_heading(doc, labels["scope_heading"], level=1, rtl=doc_rtl)
        _add_para(doc, labels["scope_body"], doc_rtl)
        scope = doc.add_table(rows=1, cols=3)
        scope.style = "Table Grid"
        scope.alignment = WD_TABLE_ALIGNMENT.RIGHT if doc_rtl else WD_TABLE_ALIGNMENT.LEFT
        for cell, text in zip(scope.rows[0].cells,
                              [labels["scope_id"], labels["scope_title"], labels["scope_state"]]):
            cell.text = text
        _style_header_row(scope, doc_rtl)
        for p in pbis:
            cells = scope.add_row().cells
            cells[0].text = str(p["id"])
            cells[1].text = p["title"]
            cells[2].text = p["state"]
        if doc_rtl:
            _set_table_rtl(scope)
            for row in scope.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        apply_rtl(para)

        # Per-PBI scenario sections
        _add_heading(doc, labels["scenarios_heading"], level=1, rtl=doc_rtl)
        total_scenarios = 0
        pbi_summary = []
        for p in pbis:
            n = _build_pbi_section(doc, labels, p, doc_rtl)
            total_scenarios += n
            pbi_summary.append({"id": p["id"], "title": p["title"], "scenarios": n})

        # Sign-off
        _add_heading(doc, labels["signoff_heading"], level=1, rtl=doc_rtl)
        _add_para(doc, labels["signoff_body"], doc_rtl)
        signoff = doc.add_table(rows=1, cols=4)
        signoff.style = "Table Grid"
        signoff.alignment = WD_TABLE_ALIGNMENT.RIGHT if doc_rtl else WD_TABLE_ALIGNMENT.LEFT
        for cell, text in zip(signoff.rows[0].cells,
                              [labels["role"], labels["name"], labels["signature"], labels["date"]]):
            cell.text = text
        _style_header_row(signoff, doc_rtl)
        for role in labels["roles"]:
            cells = signoff.add_row().cells
            cells[0].text = role
        if doc_rtl:
            _set_table_rtl(signoff)
            for row in signoff.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        apply_rtl(para)

        # ── Resolve destination & save ──────────────────────────────────────
        if not output_path.strip():
            repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            out_dir = os.path.join(repo_root, "outputs", "UAT")
            safe_sprint = re.sub(r'[^\w\-]+', '_', sprint_short).strip('_') or "sprint"
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(out_dir, f"UAT_{safe_sprint}_{stamp}.docx")

        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        doc.save(output_path)

        return {
            "status": "created",
            "file_path": os.path.abspath(output_path),
            "sprint": iteration_path,
            "language": language,
            "total_pbis": len(pbis),
            "total_scenarios": total_scenarios,
            "pbis": pbi_summary,
        }

    except Exception as e:
        return handle_error(e, "create_uat_document_from_sprint")

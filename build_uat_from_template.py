# -*- coding: utf-8 -*-
"""
Build a client UAT document by CLONING the approved template
(قالب_حالات_الاختبار_UAT.docx) and injecting real, derived UAT test cases.

Each backlog item -> one or more UAT cases. Each case becomes a clone of the
template's per-case table (UAT-XXX), with numbered steps. The two tester-result
columns are left blank for the testers to fill.

Data-driven: edit CONFIG + CASES, then run. Reusable for the full sprint.
"""
import copy
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W = qn  # shorthand

# ── helpers ──────────────────────────────────────────────────────────────────

def _set_run_text(r, text):
    """Set a run's text, clearing existing w:t but keeping rPr formatting."""
    for t in r.findall(W('w:t')):
        r.remove(t)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)


def set_tc_text(tc, text):
    """Set the text of a table cell (lxml w:tc), preserving first-run format."""
    ps = tc.findall(W('w:p'))
    p = ps[0]
    runs = p.findall(W('w:r'))
    for extra in runs[1:]:
        p.remove(extra)
    if runs:
        _set_run_text(runs[0], text)
    else:
        r = OxmlElement('w:r')
        _set_run_text(r, text)
        p.append(r)
    for extra in ps[1:]:
        tc.remove(extra)


def set_paragraph_text(p, new_text):
    """Replace a paragraph's text, keeping the first run's formatting.

    Accepts a python-docx Paragraph or an lxml w:p element."""
    el = getattr(p, '_p', p)
    runs = el.findall(W('w:r'))
    if not runs:
        r = OxmlElement('w:r'); _set_run_text(r, new_text); el.append(r); return
    _set_run_text(runs[0], new_text)
    for extra in runs[1:]:
        el.remove(extra)


def para_text(p):
    el = getattr(p, '_p', p)
    return "".join(n.text or '' for n in el.iter(W('w:t')))


def replace_in_paragraphs(doc, replacements):
    """For each paragraph, apply ALL matching replacements, then set once."""
    for p in doc.paragraphs:
        full = para_text(p)
        new = full
        for key, val in replacements.items():
            if key in new:
                new = new.replace(key, val)
        if new != full:
            set_paragraph_text(p, new)


def build_case_table(template_tbl, case, client_result_label):
    """Deep-copy the template per-case table and fill it for one case."""
    tbl = copy.deepcopy(template_tbl)
    trs = tbl.findall(W('w:tr'))
    title_tr, precond_tr, header_tr = trs[0], trs[1], trs[2]
    step_template = trs[3]

    # title (merged, single logical cell -> first w:tc)
    set_tc_text(title_tr.find(W('w:tc')),
                f"UAT-{case['no']:03d} - {case['title']}")
    # preconditions (merged)
    set_tc_text(precond_tr.find(W('w:tc')),
                f"المتطلبات الأساسية: {case['preconditions']}")
    # rename the client-side result column header (col index 4)
    if client_result_label:
        hcells = header_tr.findall(W('w:tc'))
        set_tc_text(hcells[4], client_result_label)

    # rebuild step rows
    for tr in trs[3:]:
        tbl.remove(tr)
    for i, (action, expected) in enumerate(case['steps'], start=1):
        tr = copy.deepcopy(step_template)
        cells = tr.findall(W('w:tc'))
        set_tc_text(cells[0], str(i))
        set_tc_text(cells[1], action)
        set_tc_text(cells[2], expected)
        set_tc_text(cells[3], "")
        set_tc_text(cells[4], "")
        tbl.append(tr)
    return tbl


# ── config + cases ───────────────────────────────────────────────────────────

CONFIG = {
    "template": "/Users/asmaa/CI-CD-AI-AU/قالب_حالات_الاختبار_UAT.docx",
    "output":   "/Users/asmaa/CI-CD-AI-AU/azure-mcp---Repo/outputs/UAT/"
                "UAT_CamelRace_Sprint1_SAMPLE.docx",
    "project_name": "سباقات الهجن",
    "phase":        "تطوير تطبيق الهاتف المحمول – Sprint 1",
    "client":       "اللجنة المنظمة لسباقات الهجن",   # CONFIRM/EDIT
    "client_result_label": "نتيجة اختبار العميل",      # was: إدارة المساجد
    "group_name":   "حالات اختبار قبول المستخدم – تطبيق سباقات الهجن (Sprint 1)",
    "author":       "iHorizons – فريق ضمان الجودة",
    "doc_date":     "2026/06/21",
}

CASES = [
    # ── PBI 123679 — Splash screen ──
    {"pbi": 123679, "no": 1,
     "title": "التحقق من عرض شاشة البداية بعد تسجيل الدخول",
     "preconditions": "المستخدم يملك حساباً صالحاً وقام بتسجيل الدخول إلى التطبيق.",
     "steps": [
        ("سجّل الدخول إلى التطبيق ببيانات صحيحة.",
         "تظهر شاشة البداية (Splash) مباشرة بعد تسجيل الدخول."),
        ("راقب شاشة البداية حتى اكتمال تحميل الصفحة الرئيسية.",
         "تبقى شاشة البداية ظاهرة إلى أن يتم فتح الصفحة الرئيسية للتطبيق."),
        ("تحقق من مطابقة شاشة البداية للتصميم المعتمد.",
         "تظهر شاشة البداية بالشعار والألوان والعناصر وفق التصميم المعتمد."),
     ]},

    # ── PBI 125308 — السباقات ──
    {"pbi": 125308, "no": 2,
     "title": "التحقق من صفحة السباقات والعرض الافتراضي",
     "preconditions": "المستخدم مسجّل الدخول ويوجد سباقات مرتبطة بالموسم الحالي.",
     "steps": [
        ("من القائمة أو الواجهة الرئيسية اضغط على خيار «السباقات».",
         "يتم الانتقال إلى صفحة السباقات بنجاح."),
        ("لاحظ أعلى الصفحة.",
         "يظهر بانر تعريفي يتيح الوصول المباشر إلى الكتيب الخاص بالسباقات."),
        ("لاحظ قائمة السباقات المعروضة افتراضياً.",
         "تُعرض جميع سباقات الموسم الحالي (أو آخر موسم متاح) مع اسم السباق واسم "
         "المهرجان وتاريخ البداية وتاريخ الانتهاء."),
     ]},
    {"pbi": 125308, "no": 3,
     "title": "التحقق من البحث في السباقات",
     "preconditions": "المستخدم متواجد في صفحة السباقات.",
     "steps": [
        ("أدخل اسم سباق موجود في حقل البحث.",
         "تُعرض السباقات المطابقة لاسم السباق المُدخل."),
        ("امسح البحث وأدخل اسم مهرجان موجود.",
         "تُعرض السباقات المطابقة لاسم المهرجان المُدخل."),
        ("أدخل قيمة غير موجودة في حقل البحث.",
         "تظهر حالة «لا توجد نتائج» دون أي خطأ في النظام."),
     ]},
    {"pbi": 125308, "no": 4,
     "title": "التحقق من الفلترة حسب الموسم والنوع والمرحلة ودمج الفلاتر",
     "preconditions": "المستخدم متواجد في صفحة السباقات.",
     "steps": [
        ("غيّر الموسم واختر موسماً آخر.",
         "تُعرض جميع السباقات المرتبطة بالموسم المحدد."),
        ("طبّق الفلترة حسب نوع السباق (سباق / مهرجان).",
         "تُعرض السباقات التابعة للنوع المحدد فقط."),
        ("طبّق الفلترة حسب المرحلة (الأولى / الثانية).",
         "تُعرض السباقات التابعة للمرحلة المحددة فقط."),
        ("طبّق أكثر من فلتر في آنٍ واحد (موسم + نوع + مرحلة).",
         "تتحدث قائمة السباقات لتعرض النتائج المطابقة لجميع المعايير مجتمعة."),
     ]},

    # ── PBI 125313 — تفاصيل المطية ──
    {"pbi": 125313, "no": 5,
     "title": "التحقق من عرض تفاصيل المطية ومعلوماتها الأساسية",
     "preconditions": "المستخدم مسجّل الدخول ويوجد مطية في قائمة الهجن.",
     "steps": [
        ("من قائمة الهجن اضغط على اسم أي مطية.",
         "يتم الانتقال إلى صفحة تفاصيل المطية الخاصة بها."),
        ("لاحظ أعلى الشاشة.",
         "تُعرض صورة المطية بشكل واضح في أعلى الصفحة."),
        ("تحقق من المعلومات الأساسية أسفل الصورة.",
         "يُعرض اسم المطية وحالتها باللون المخصص لها والعمر والجنس ورقم الشريحة "
         "الإلكترونية."),
        ("تحقق من معلومات الأب والأم.",
         "يُعرض اسم الأب واسم الأم، وعند عدم توفر المعلومة تظهر علامة (—)."),
        ("تحقق من الفئة ونوع الملكية والمضمر.",
         "تُعرض فئة المطية ونوع ملكيتها واسم المضمر، وعند عدم وجود مضمر يُعرض اسم "
         "المالك بدلاً منه."),
     ]},
    {"pbi": 125313, "no": 6,
     "title": "التحقق من نسخ رقم الشريحة وإعداد إظهار المطية للعامة",
     "preconditions": "المستخدم متواجد في صفحة تفاصيل المطية وهو مالك المطية.",
     "steps": [
        ("اضغط على زر نسخ رقم الشريحة.",
         "يتم نسخ رقم الشريحة إلى الحافظة (Clipboard) مع تأكيد عملية النسخ."),
        ("فعّل خيار «إظهار معلومات المطية للعامة».",
         "تصبح معلومات المطية ظاهرة للزوار."),
        ("ألغِ تفعيل الخيار لإخفاء المطية عن الزوار.",
         "تُخفى معلومات المطية عن الزوار."),
     ]},
]


# ── build ────────────────────────────────────────────────────────────────────

def main():
    doc = Document(CONFIG["template"])

    # header / intro / group-name replacements
    replace_in_paragraphs(doc, {
        "[اسم المشروع]": CONFIG["project_name"],
        "[المرحلة]": CONFIG["phase"],
        "وزارة الأوقاف والشؤون الإسلامية": CONFIG["client"],
        "[اسم المشروع / المرحلة]": f"{CONFIG['project_name']} – {CONFIG['phase']}",
        "[اسم مجموعة حالات الاختبار]": CONFIG["group_name"],
    })

    # version-control table (doc.tables[1]) row 1
    vt = doc.tables[1]
    vt.rows[1].cells[0].text = "1.0"
    vt.rows[1].cells[1].text = CONFIG["doc_date"]
    vt.rows[1].cells[2].text = CONFIG["author"]
    vt.rows[1].cells[3].text = "إنشاء الملف"

    # capture template tables: T2 editor-note, T3/T4 examples, T5 placeholder
    note_tbl   = doc.tables[2]._tbl
    example_a  = doc.tables[3]._tbl
    example_b  = doc.tables[4]._tbl
    case_tmpl  = doc.tables[5]._tbl

    # insert generated case tables before the editor-note table, with spacers
    for case in CASES:
        new_tbl = build_case_table(case_tmpl, case, CONFIG["client_result_label"])
        note_tbl.addprevious(new_tbl)
        spacer = OxmlElement('w:p')
        note_tbl.addprevious(spacer)

    # remove the template tables (note + 2 examples + placeholder)
    for el in (note_tbl, example_a, example_b, case_tmpl):
        el.getparent().remove(el)

    doc.save(CONFIG["output"])
    print("Saved:", CONFIG["output"])
    print("Cases:", len(CASES),
          "across PBIs:", sorted({c['pbi'] for c in CASES}))


if __name__ == "__main__":
    main()
